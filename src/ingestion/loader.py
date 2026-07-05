"""
src/ingestion/loader.py
Document loading engine supporting PDF, DOCX, and TXT.

Design decisions:
- Each loader returns a list of Document objects with text + metadata.
- PDF loader extracts text page-by-page so we preserve page numbers.
- DOCX loader extracts paragraph-by-paragraph to preserve structure.
- TXT loader reads the entire file as one document.
- All loaders are defensive: corrupted files return empty list with a logged warning
  rather than crashing the pipeline.
- Optional imports (PyPDF2, python-docx) are handled gracefully so the app can
  start even if one dependency is temporarily missing.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Union

from loguru import logger

# Optional imports handled gracefully
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


@dataclass
class Document:
    """
    Internal data contract for a loaded document page, paragraph, or file.
    Travels through the entire pipeline: loader -> chunker -> embedder -> vector store.
    """
    text: str
    source: str          # Original file path
    metadata: dict       # e.g., {"page": 3, "type": "pdf", "chunk_index": 0}


class DocumentLoader:
    """
    Unified entrypoint for loading documents from multiple formats.
    Dispatches to the correct parser based on file extension.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def load_file(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a single file and return a list of Document objects.
        Returns empty list on unsupported format, missing file, or I/O error.
        """
        path = Path(file_path).resolve()
        if not path.exists():
            logger.error("File not found: {}", path)
            return []

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.error("Unsupported file format: {} | path={}", ext, path)
            return []

        logger.info("Loading document | path={} | format={}", path, ext)

        try:
            if ext == ".pdf":
                return self._load_pdf(path)
            elif ext == ".docx":
                return self._load_docx(path)
            elif ext == ".txt":
                return self._load_txt(path)
        except Exception as e:
            logger.exception("Failed to load document | path={} | error={}", path, e)
            return []

        return []  # Defensive fallback

    def load_directory(self, dir_path: Union[str, Path]) -> List[Document]:
        """
        Recursively load all supported documents from a directory.
        """
        path = Path(dir_path).resolve()
        if not path.is_dir():
            logger.error("Not a directory: {}", path)
            return []

        documents: List[Document] = []
        for file_path in path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                docs = self.load_file(file_path)
                documents.extend(docs)

        logger.info(
            "Directory ingestion complete | dir={} | docs={}", path, len(documents)
        )
        return documents

    # ------------------------------------------------------------------
    # Format-specific loaders
    # ------------------------------------------------------------------
    def _load_pdf(self, path: Path) -> List[Document]:
        if PdfReader is None:
            logger.error("PyPDF2 not installed. Cannot load PDF: {}", path)
            return []

        reader = PdfReader(str(path))
        documents: List[Document] = []

        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text:
                documents.append(
                    Document(
                        text=text,
                        source=str(path),
                        metadata={"page": i, "type": "pdf"},
                    )
                )
        logger.debug("PDF loaded | pages={} | path={}", len(documents), path)
        return documents

    def _load_docx(self, path: Path) -> List[Document]:
        if DocxDocument is None:
            logger.error("python-docx not installed. Cannot load DOCX: {}", path)
            return []

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # For DOCX we treat the whole file as one document for now.
        # Paragraph-level splitting happens in the chunker.
        return [
            Document(
                text=full_text,
                source=str(path),
                metadata={"type": "docx", "paragraphs": len(paragraphs)},
            )
        ]

    def _load_txt(self, path: Path) -> List[Document]:
        text = path.read_text(encoding="utf-8")
        return [
            Document(
                text=text,
                source=str(path),
                metadata={"type": "txt"},
            )
        ]